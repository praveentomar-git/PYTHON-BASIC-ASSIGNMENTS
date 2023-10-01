#!/usr/bin/env python
# coding: utf-8

# ## Q1. In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened?

# ### ANS1. These files will be opened in binary mode, read binary (rb) for PdfFileReader() and write binary (wb) PdfFileWriter() but these classes have been changed in the new version of pyPDF2.

# ## Q2. From a PdfFileReader object, how do you get a Page object for page 5?

# ### ANS2. Calling pages(4) will return a Page object for page 5 since page 0 is the first page.

# In[2]:


#!pip install PyPDF2


# In[10]:


import PyPDF2 as pdf
pdfFileObj = open("KT.pdf",'rb')
pdfReader = pdf.PdfReader(pdfFileObj)
pageObj = pdfReader.pages[4]
pageObj.extract_text()


# ## Q3. What PdfFileReader variable stores the number of pages in the PDF document?

# In[6]:


import PyPDF2 as pdf
pdfFileObj = open("KT.pdf",'rb')
pdfReader = pdf.PdfReader(pdfFileObj)
len(pdfReader.pages)    # Prints the no. of pages in an input document.


# ## Q4. If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do before you can obtain Page objects from it?

# ### ANS4. Before we obtain the page object, the pdf has to be decrypted by calling .decrypt('swordfish')

# In[8]:


#pdfReader.decrypt('swordfish')


# ## Q5. What methods do you use to rotate a page?

# ### ANS5. The rotate() method. The degree to rotate is passed as an integer argument.

# In[15]:


pageObj.rotate(360)


# ## Q6. What is the difference between a Run object and a Paragraph object?

# ### ANS6. A Run object represents a formatted portion of text within a paragraph, defining its style, font, and other attributes. It is the smallest text unit in a document.
# 
# ### A Paragraph object represents a larger block of text separated by line or paragraph breaks, with its own formatting settings, such as alignment and indentation. It organizes content into coherent sections.

# ## Q7. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc?

# In[25]:


#!pip install python-docx
import docx
doc = docx.Document('CourseOutline.docx')
doc.paragraphs


# In[26]:


for paragraph in doc.paragraphs:
    print(paragraph.text)


# ## Q8. What type of object has bold, underline, italic, strike, and outline variables?

# ### ANS8. A Run object has bold, underline,italic,strike and outline variables.

# ## Q9. What is the difference between False, True, and None for the bold variable?

# ### ANS9. In the context of a bold variable:
# 
# `False:` Indicates that the text should not be bold. The text remains in its normal, non-bold style.
# 
# `True:` Indicates that the text should be bold. The text is displayed with a bold style.
# 
# `None:` Typically means that the bold formatting is not explicitly specified. The text takes on the default formatting, which may or may not be bold, depending on the context or the default settings of the software or system.
bold = True  # Style Set to Bold
bold = False # Style Not Set to Bold
bold = None # Style is Not Applicable
# ## Q10. How do you create a Document object for a new Word document?

# In[27]:


from docx import Document
document = Document()
document.add_paragraph("iNeuron's Full Stack DataScience Course")
document.save('mydocument.docx')


# ## Q11. How do you add a paragraph with the text &#39;Hello, there!&#39; to a Document object stored in a variable named doc?

# In[29]:


import docx
doc = docx.Document()
doc.add_paragraph('Hello there!')
doc.save('hellothere.docx')


# ## Q12. What integers represent the levels of headings available in Word documents?

# ### ANS12. The integers that represent the levels of headings available in Word documents in Python docx are 0 to 9, inclusive.
# 
# * 0: Title
# * 1: Heading 1
# * 2: Heading 2
# * 3: Heading 3
# * 4: Heading 4
# * 5: Heading 5
# * 6: Heading 6
# * 7: Heading 7
# * 8: Heading 8
# * 9: Heading 9
# 
# We can use these integers to add headings to our Word documents using the add_heading() method on the Document object.
